#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
Created on Thu Sep 22 14:09:55 2022
"""
import fitz
import os
import docx
from PIL import Image
import pdfplumber
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
import aspose.words as aw
import pytesseract
import PyPDF2
import shutil
import re
import docx2txt
import logging
logging.getLogger("pdfminer").setLevel(logging.WARNING)
logging.getLogger("PyPDF2").setLevel(logging.WARNING)
logging.getLogger("docx2txt").setLevel(logging.WARNING)
logging.getLogger("pdfplumber").setLevel(logging.WARNING)


pytesseract.pytesseract.tesseract_cmd =r'C:/Program Files/Tesseract-OCR/tesseract.exe'

class Extract:
    def __init__(self):
        pass
    """
    Class to retrieve plain text in raw format from documents(.pdf/.doc formats)
    using different tools and libraries based on specific requirements
    """
    @classmethod
    def ex_pdf_plumber(self, file_path):
        """
        Process files(.pdf) format using python 'pdfplumber' to parse text.
        params:
            filepath->str: path to pdf file
        returns:
            -text_all->str: Extracted text
            -None-> object: None incase of any exception/error
        """
        try:
            text_all = ''
            with pdfplumber.open(file_path) as pdf:
                pages = pdf.pages
                try:
                    pages = [pages[0], pages[1], pages[2], pages[-1]]
                except:
                    pass
                for page in pages:
                    text = page.extract_text(x_tolerance=5).strip()
                    for line in text.splitlines():
                        if line.strip():
                            text_all = text_all+line+'\n'
                        else:
                            continue
            if text_all.strip():
                return text_all.strip()
            else:
                return None
        except TypeError:
            return None
        except AttributeError:
            return None
        finally:
            return None
    
    @classmethod
    def ex_pdfminer(self, file_path):
        """
        Process files(.pdf) format using python 'pdfminer' to parse text.
        params:
            filepath->str: path to pdf file
        returns:
            -text_all->str: Extracted text
            -None-> object: None incase of any exception/error
        """
        text_all = ''
        pages = list(extract_pages(file_path))
        try:
            pages = [pages[0], pages[1], pages[2], pages[-1]]
        except:
            pass
        for page_layout in pages:
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    txt = element.get_text()
                    if txt.strip():
                        text_all = text_all+txt.strip()+'\n'
                    else:
                        continue
        if text_all.strip():
            text_all = '\n'.join([x for x in text_all.split('/n')])
            return text_all.strip()
        else:
            return None

    
    @classmethod
    def ex_pypdf(self, file_path):
        """
        Process files(.pdf) format using python 'pyPDF2' to parse text.
        params:
            filepath->str: path to pdf file
        returns:
            -text->str: Extracted text
            -None-> object: None incase of any exception/error
        """

        try:
            pdfFileObj = open(file_path, 'rb')	
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
            n = pdfReader.numPages
            x = 0
            text = ''
            for page in pdfReader.pages:
                x = x+1
                if x<4 or x==n:
                    for line in page.extractText().splitlines():
                        if line.strip():
                            text = text + line + '\n'
                        else:
                            continue
                else:
                    break
            if text.strip():
                return text.strip()
            else:
                return None
        except:
            return None
    
    @classmethod
    def ocr_text(self, file_path):
        """
        Process files(.pdf) format using python 'tesserocr' to parse text.
        params:
            filepath->str: path to pdf file
        returns:
            -text_all->str: Extracted text
            -None-> object: None incase of any exception/error
        """
        doc = fitz.open(file_path)
        n = doc.page_count-1
        file_name = file_path.split('\\')[-1]
        temp_path = 'temp\\{fn}'.format(fn=file_name)
        pages = [0,1,2,n]
            
        for i in pages:
            try:
                if not os.path.exists(temp_path):
                    os.mkdir(temp_path)
                page = doc.load_page(i)
                pix = page.get_pixmap()
                output = temp_path + '/' + str(i)+'.png'
                pix.save(output)
            except ValueError:
                pass
                
        text_all = ''
        filelist = os.listdir(temp_path+'/')
        filelist = sorted(filelist,key=lambda x: int(os.path.splitext(x)[0])) 
        for index, name in enumerate(filelist):
            pth = temp_path +'/'+name
            img = Image.open(pth)
            new_size = tuple(4*x for x in img.size)
            img = img.resize(new_size, Image.ANTIALIAS)
            text = pytesseract.image_to_string(img, lang='spa', config='--psm 4 --psm 6')
            for line in text.strip().splitlines():
                line = line.strip()
                if line:
                    text_all = text_all + line + '\n'
        shutil.rmtree(temp_path)
        if text_all.strip():
            return text_all.strip()
        else:
            return None
        
        
    @classmethod
    def extract_doc(self, file_path):
        """
        Process files(.doc/.docx) format using python 'docx' to parse text.
        params:
            filepath->str: path to pdf file
        returns:
            -text_all->str: Extracted text
            -None-> object: None incase of any exception/error
        """

        doc_path = {"path": file_path, 'status': "orig"}
        if doc_path['path'].endswith('.doc'):
            file_path = "C:/Users/engr2/OneDrive/Desktop/Legal-Information-Retrieval/Legis documents/prj_809567_S2001-01402CSJ.doc"
            doc = aw.Document(doc_path['path'])
            doc_path['path'] = doc_path['path']+'x'
            if doc.has_macros:
                doc.remove_macros()
            else:
                pass
            doc.save(doc_path['path'])
            doc_path['status'] = 'changed'
        else:
            pass

        try:
            text_all = docx2txt.process(doc_path['path'])
            text_all = '\n'.join([x.strip() for x in text_all.split('\n') if x.strip()!=''])
        except:
            document = docx.Document(doc_path['path'])
            doc_par = (
                paragraph.text for paragraph in document.paragraphs
            )
            text_all = ''
            for index, par in enumerate(doc_par):
                par = par.strip()
                if par:
                    for line in par.splitlines():
                        text_all = text_all + line.strip() + '\n'
                else:
                    continue
        if doc_path['status']=='changed':
            os.remove(doc_path['path'])        
        else:
            pass
        
        if text_all:    
            return text_all.strip()
        else:
            return None
    
    def get(self, file_path):
        """
        Comfile all extraction functions into on to process the input files
        for text extractioin.
        params:
            file_path-> str: input path to file(doc/docx/pdf)
        returns:
            text-> str: extracted text
        """
        if file_path.endswith('.doc') or file_path.endswith('.docx'):
            text = self.extract_doc(file_path=file_path)
        elif file_path.endswith('.pdf'):
            text = self.ex_pdf_plumber(file_path=file_path)
            if text and re.findall('[\u4e00-\u9fff]', text)==[]:
                text = text
            else:
                text = self.ex_pdfminer(file_path=file_path)
                if text and re.findall('[\u4e00-\u9fff|]', text)==[]:
                    text = text
                else:
                    text = self.ocr_text(file_path=file_path)
                    if text:
                        text = text
                    else:
                        text = ''
        else:
            text = ''
        return text
